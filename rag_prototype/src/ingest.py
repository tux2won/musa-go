from __future__ import annotations

import csv
import hashlib
import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

import fitz
from docx import Document

from .config import (
    CHUNKS_PATH,
    DISEASE_ALIASES,
    DOCUMENTS_PATH,
    INTENT_KEYWORDS,
    MANIFEST_PATH,
    SOURCE_DIR,
    ensure_directories,
)


@dataclass
class SourceDocument:
    document_id: str
    title: str
    source_path: str
    source_org: str
    source_type: str
    authority_level: str
    effective_status: str
    effective_date: str
    retrieved_at: str
    checksum: str
    file_type: str
    page_count: int
    char_count: int
    replacement_chars: int
    use_for_rag: bool
    exclusion_reason: str


@dataclass
class Chunk:
    chunk_id: str
    document_id: str
    title: str
    source_path: str
    source_org: str
    source_type: str
    authority_level: str
    effective_status: str
    effective_date: str
    locator: str
    page: int | None
    section_path: str
    disease_tags: list[str]
    intent_tags: list[str]
    text: str
    coverage_status: str


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    text = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\(#AJAX\)", r"\1", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def infer_source(path: Path) -> tuple[str, str, str, str, str]:
    name = path.name
    lowered = str(path).lower()
    if "나라사랑" in name or "kb_" in lowered:
        return "KB국민은행·KB손해보험(제공 이미지 정리)", "insurance_terms", "D", "historical_unverified", "unknown"
    if "국방" in name or "군보건" in name or "공무상" in name or "민간위탁" in name or "청원휴가" in name:
        return "국방부·국가법령정보센터", "directive" if "훈령" in name else "law", "A", "current", "2026"
    if "건강보험심사평가원" in name or "본인부담" in name or "진료비 확인" in name:
        return "건강보험심사평가원", "public_guidance", "B", "current", "2026"
    if "국민건강보험" in name:
        return "찾기쉬운 생활법령정보", "public_guidance", "B", "current", "2026-05-15"
    if "법령" in name or "감염병의 예방 및 관리에 관한 법률" in name:
        return "국가법령정보센터", "law", "A", "current", "2026"
    return "질병관리청", "guideline" if path.suffix.lower() == ".pdf" else "official_web", "A", "current", "2026"


def document_id(path: Path) -> str:
    rel = path.relative_to(SOURCE_DIR).as_posix()
    return hashlib.sha1(rel.encode("utf-8")).hexdigest()[:14]


def detect_tags(text: str) -> tuple[list[str], list[str]]:
    low = text.lower()
    diseases = [
        disease
        for disease, aliases in DISEASE_ALIASES.items()
        if any(alias.lower() in low for alias in aliases)
    ]
    intents = [
        intent
        for intent, keywords in INTENT_KEYWORDS.items()
        if any(keyword.lower() in low for keyword in keywords)
    ]
    return diseases, intents


def split_long(text: str, max_chars: int = 1800, overlap_chars: int = 180) -> list[str]:
    text = clean_text(text)
    if len(text) <= max_chars:
        return [text] if len(text) >= 40 else []
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    parts: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            sentences = re.split(r"(?<=[.!?다요함됨])\s+", paragraph)
        else:
            sentences = [paragraph]
        for sentence in sentences:
            candidate = f"{current}\n\n{sentence}".strip()
            if current and len(candidate) > max_chars:
                parts.append(current)
                current = (current[-overlap_chars:] + " " + sentence).strip()
            else:
                current = candidate
    if current:
        parts.append(current)
    return [part for part in parts if len(part) >= 40]


def markdown_sections(text: str) -> Iterable[tuple[str, str]]:
    current_heading = "문서 개요"
    buffer: list[str] = []
    for line in text.splitlines():
        match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
        article = re.match(r"^(제\d+조(?:의\d+)?\([^)]*\))", line.strip())
        if (match or article) and buffer:
            yield current_heading, "\n".join(buffer)
            buffer = []
        if match:
            current_heading = match.group(2).strip()
        elif article:
            current_heading = article.group(1).strip()
            buffer.append(line)
        else:
            buffer.append(line)
    if buffer:
        yield current_heading, "\n".join(buffer)


def extract_units(path: Path) -> tuple[list[tuple[str, int | None, str]], int]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        doc = fitz.open(path)
        units = [(f"PDF p.{index + 1}", index + 1, page.get_text("text")) for index, page in enumerate(doc)]
        return units, len(doc)
    if suffix == ".md":
        text = path.read_text(encoding="utf-8", errors="replace")
        return [(heading, None, body) for heading, body in markdown_sections(text)], 0
    if suffix == ".csv":
        for encoding in ("utf-8-sig", "cp949"):
            try:
                with path.open(encoding=encoding, newline="") as handle:
                    rows = list(csv.DictReader(handle))
                break
            except UnicodeDecodeError:
                continue
        units = []
        for i, row in enumerate(rows, start=1):
            term = row.get("단어명") or row.get("용어") or f"행 {i}"
            body = " | ".join(f"{key}: {value}" for key, value in row.items() if value)
            units.append((str(term), None, body))
        return units, 0
    if suffix == ".docx":
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return list(markdown_sections(text)), 0
    return [], 0


def should_use(path: Path) -> tuple[bool, str]:
    name = path.name
    suffix = path.suffix.lower()
    if suffix == ".doc":
        return False, "동일 훈령의 Markdown 정리본 사용"
    if suffix == ".docx" and "나라사랑" in name:
        return False, "동일 내용 Markdown 정리본 사용"
    if suffix == ".csv" and "용어사전" in name:
        return False, "동일 내용 RAG용 Markdown 정리본 사용"
    if name == "A형간염_질병관리청_국가건강정보포털.md":
        return False, "동일 내용 중복 파일"
    return suffix in {".md", ".pdf", ".csv", ".docx"}, ""


def build_corpus() -> tuple[list[SourceDocument], list[Chunk]]:
    ensure_directories()
    documents: list[SourceDocument] = []
    chunks: list[Chunk] = []
    for path in sorted(SOURCE_DIR.rglob("*")):
        if not path.is_file():
            continue
        use_for_rag, reason = should_use(path)
        source_org, source_type, authority, effective_status, effective_date = infer_source(path)
        units: list[tuple[str, int | None, str]] = []
        page_count = 0
        if use_for_rag:
            try:
                units, page_count = extract_units(path)
            except Exception as exc:
                use_for_rag = False
                reason = f"추출 실패: {type(exc).__name__}"
        total_text = "\n".join(unit[2] for unit in units)
        doc_id = document_id(path)
        record = SourceDocument(
            document_id=doc_id,
            title=path.stem,
            source_path=str(path),
            source_org=source_org,
            source_type=source_type,
            authority_level=authority,
            effective_status=effective_status,
            effective_date=effective_date,
            retrieved_at="2026-07-02",
            checksum=sha256(path),
            file_type=path.suffix.lower().lstrip("."),
            page_count=page_count,
            char_count=len(total_text),
            replacement_chars=total_text.count("\ufffd"),
            use_for_rag=use_for_rag,
            exclusion_reason=reason,
        )
        documents.append(record)
        if not use_for_rag:
            continue
        for unit_index, (locator, page, raw_text) in enumerate(units, start=1):
            cleaned = clean_text(raw_text)
            for part_index, part in enumerate(split_long(cleaned), start=1):
                diseases, intents = detect_tags(f"{path.stem} {locator} {part}")
                chunk_id = f"{doc_id}_{unit_index:04d}_{part_index:02d}"
                chunks.append(
                    Chunk(
                        chunk_id=chunk_id,
                        document_id=doc_id,
                        title=path.stem,
                        source_path=str(path),
                        source_org=source_org,
                        source_type=source_type,
                        authority_level=authority,
                        effective_status=effective_status,
                        effective_date=effective_date,
                        locator=locator,
                        page=page,
                        section_path=locator,
                        disease_tags=diseases,
                        intent_tags=intents,
                        text=part,
                        coverage_status="unknown" if source_type == "insurance_terms" else "not_applicable",
                    )
                )
    return documents, chunks


def save_corpus(documents: list[SourceDocument], chunks: list[Chunk]) -> None:
    ensure_directories()
    with DOCUMENTS_PATH.open("w", encoding="utf-8") as handle:
        for document in documents:
            handle.write(json.dumps(asdict(document), ensure_ascii=False) + "\n")
    with CHUNKS_PATH.open("w", encoding="utf-8") as handle:
        for chunk in chunks:
            handle.write(json.dumps(asdict(chunk), ensure_ascii=False) + "\n")
    with MANIFEST_PATH.open("w", encoding="utf-8-sig", newline="") as handle:
        rows = [asdict(document) for document in documents]
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def main() -> None:
    documents, chunks = build_corpus()
    save_corpus(documents, chunks)
    print(f"documents={len(documents)} indexed={sum(d.use_for_rag for d in documents)} chunks={len(chunks)}")
    print(f"replacement_chars={sum(d.replacement_chars for d in documents)}")


if __name__ == "__main__":
    main()
